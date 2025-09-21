from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

def set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_full_row_top_border(row):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        tcPr.append(tcBorders)

# --- Load Data ---
df = pd.read_csv(r"Results/r18/summary_metrics.csv")

# Round numeric metrics to 3 decimals
metrics_cols = ["std_ref","std_model","rmse","corr","crmse","mae","mape"]
df[metrics_cols] = df[metrics_cols].applymap(
    lambda x: round(x, 3) if isinstance(x, (int, float)) else x
)

# Force timescale ordering
timescale_order = ["SPI_1","SPI_3","SPI_6","SPI_9","SPI_12","SPI_24"]
df["spi"] = pd.Categorical(df["spi"], categories=timescale_order, ordered=True)

# --- Create Document ---
doc = Document()

cols = ["Station", "Timescale", "Std Ref", "Model","Std Model","RMSE","Corr","CRMSE","MAE","MAPE"]
cell_widths = [2000, 2000, 1500, 2000, 2000, 1200, 1200, 1200, 1200, 1200]

# Add header row
table = doc.add_table(rows=1, cols=len(cols))
table.style = None
hdr_cells = table.rows[0].cells
for i, col in enumerate(cols):
    hdr_cells[i].text = col
    set_cell_width(hdr_cells[i], cell_widths[i])

# --- Fill Table ---
for station, station_group in df.groupby("station"):
    station_start_idx = len(table.rows)

    for spi, spi_group in station_group.groupby("spi", sort=False):
        timescale_start_idx = len(table.rows)

        for idx, (_, r) in enumerate(spi_group.iterrows()):
            row = table.add_row().cells
            # only first row gets station and timescale labels
            if idx == 0:
                row[0].text = str(station)
                row[1].text = str(spi)
            row[2].text = str(r["std_ref"])
            for j, col in enumerate(["model","std_model","rmse","corr","crmse","mae","mape"], start=3):
                row[j].text = str(r[col])
                row[j].paragraphs[0].alignment = 2  # right align numbers
            for k, cell in enumerate(row):
                set_cell_width(cell, cell_widths[k])
                remove_all_borders(cell)

        # merge Timescale column for all rows in this block
        end_idx = len(table.rows)-1
        if end_idx > timescale_start_idx:
            table.cell(timescale_start_idx,1).merge(table.cell(end_idx,1))

        # add top border for whole block
        add_full_row_top_border(table.rows[timescale_start_idx])

    # merge Station column for whole group
    end_station_idx = len(table.rows)-1
    if end_station_idx > station_start_idx:
        table.cell(station_start_idx,0).merge(table.cell(end_station_idx,0))

    # add top border for station group
    add_full_row_top_border(table.rows[station_start_idx])

# --- Save ---
doc.save("Results/table_results_clean.docx")
print("Created Results/table_results_clean.docx")
