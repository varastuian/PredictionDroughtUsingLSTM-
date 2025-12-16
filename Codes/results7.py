from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx.shared import RGBColor
from docx.shared import Pt
import numpy as np
def set_cell_font(cell, size_pt=8, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)
            run.font.bold = bold

def set_cell_text_green(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 176, 80)  # Word green

def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def remove_all_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_full_row_top_border(row):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        tcPr.append(tcBorders)

# --- Load Data ---
df = pd.read_csv(r"Results/12151716_e170-hd64-l2-d0.1-h3/summary_metrics.csv")
def normalize_and_round(x):
    # string like "[0.05888152]"
    if isinstance(x, str):
        x = x.strip()
        if x.startswith('[') and x.endswith(']'):
            x = x[1:-1]

    # list / ndarray
    if isinstance(x, (list, np.ndarray)):
        x = x[0]

    # numpy scalar
    if isinstance(x, np.generic):
        x = x.item()

    try:
        return round(float(x), 1)
    except (TypeError, ValueError):
        return np.nan
metrics_cols = ["std_ref","std_model","rmse","corr","crmse","mape","score"]
df[metrics_cols] = df[metrics_cols].applymap(normalize_and_round)



# Timescale order   
timescale_order = ["SPI_1","SPI_3","SPI_6","SPI_9","SPI_12","SPI_24"]
df["spi"] = pd.Categorical(df["spi"], categories=timescale_order, ordered=True)

# --- Create Document ---
doc = Document()

cols = ["Station", "Timescale", "Std Ref", "Model","Std Model","RMSE","Corr","CRMSE","Score"]
cell_widths = [2030, 2050, 2500, 2000, 2000, 1200, 1200, 1200, 1200, 1200]

# Add header
table = doc.add_table(rows=1, cols=len(cols))
table.style = None
for i, col in enumerate(cols):
    table.rows[0].cells[i].text = col
    set_cell_width(table.rows[0].cells[i], cell_widths[i])
    header_cell = table.rows[0].cells[i]
    header_cell.text = col
    set_cell_width(header_cell, cell_widths[i])
    set_cell_font(header_cell, size_pt=8.5, bold=True)

# --- Fill Table ---
for station, station_group in df.groupby("station"):
    station_start_idx = len(table.rows)

    for spi, spi_group in station_group.groupby("spi", sort=False):
        timescale_start_idx = len(table.rows)
        stdref_start_idx = None
        min_score = spi_group["score"].min()

        for idx, (_, r) in enumerate(spi_group.iterrows()):
            row = table.add_row().cells
            if idx == 0:
                # row[1].text = str(spi)       
                row[1].text = str(int(spi.split("_")[1]))

                row[2].text = str(r["std_ref"]) # Std Ref only once
                stdref_start_idx = len(table.rows)-1
            row[3].text = str(r["model"])
            for j, col in enumerate(["std_model","rmse","corr","crmse","score"], start=4):
                row[j].text = str(r[col])
                row[j].paragraphs[0].alignment = 2
            # --- Highlight lowest score model ---
            if r["score"] == min_score:
                set_cell_text_green(row[3])

            for k, cell in enumerate(row):
                set_cell_width(cell, cell_widths[k])
                # set_cell_font(cell, size_pt=8)   # <<< key change
                remove_all_borders(cell)

        end_idx = len(table.rows)-1
        # merge Timescale
        if end_idx > timescale_start_idx:
            table.cell(timescale_start_idx,1).merge(table.cell(end_idx,1))
        # merge Std Ref
        if end_idx > stdref_start_idx:
            table.cell(stdref_start_idx,2).merge(table.cell(end_idx,2))

        # add top border across timescale block
        add_full_row_top_border(table.rows[timescale_start_idx])

    end_station_idx = len(table.rows)-1
    table.cell(station_start_idx,0).text = str(station)
    if end_station_idx > station_start_idx:
        table.cell(station_start_idx,0).merge(table.cell(end_station_idx,0))
    add_full_row_top_border(table.rows[station_start_idx])

# --- Save ---
doc.save("Results/table_results_clean.docx")
print("✅ done")
