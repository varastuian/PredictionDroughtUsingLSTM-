from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top','left','bottom','right','insideH','insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_width(cell, width):
    """Set fixed column width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# --- Load Data ---
df = pd.read_csv(r"Results\r18\summary_metrics.csv")

# Round numeric metrics to 3 decimals
metrics_cols = ["std_ref","std_model","rmse","corr","crmse","mae","mape"]
df[metrics_cols] = df[metrics_cols].applymap(lambda x: round(x, 3) if isinstance(x, (int, float)) else x)

# --- Create Document ---
doc = Document()

# Define table columns
cols = ["Station", "Timescale", "Std Ref", "Model","Std Model","RMSE","Corr","CRMSE","MAE","MAPE"]
cell_widths = [2000, 2000, 1500, 2000, 2000, 1200, 1200, 1200, 1200, 1200]

# Add header row
table = doc.add_table(rows=1, cols=len(cols))
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, col in enumerate(cols):
    hdr_cells[i].text = col
    set_cell_width(hdr_cells[i], cell_widths[i])

# --- Fill Table ---
for station, station_group in df.groupby("station"):
    station_start_idx = len(table.rows)

    for spi, spi_group in station_group.groupby("spi"):
        timescale_start_idx = len(table.rows)

        for _, r in spi_group.iterrows():
            row = table.add_row().cells
            row[0].text = str(station)      # merge later
            row[1].text = str(spi)          # merge later
            row[2].text = str(r["std_ref"])
            for i, col in enumerate(["model","std_model","rmse","corr","crmse","mae","mape"], start=3):
                row[i].text = str(r[col])
                row[i].paragraphs[0].alignment = 2  # right align numbers
            for idx, cell in enumerate(row):
                set_cell_width(cell, cell_widths[idx])

        # merge Timescale column
        end_idx = len(table.rows)-1
        if end_idx > timescale_start_idx:
            table.cell(timescale_start_idx,1).merge(table.cell(end_idx,1))

    # merge Station column
    end_station_idx = len(table.rows)-1
    if end_station_idx > station_start_idx:
        table.cell(station_start_idx,0).merge(table.cell(end_station_idx,0))

    # add horizontal line separator after each station
    sep_row = table.add_row().cells
    sep_row[0].merge(sep_row[-1])
    tc = sep_row[0]._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tcBorders.append(bottom)
    tcPr.append(tcBorders)

# --- Save ---
doc.save("Results/table_results.docx")
print("Created Results/table_results.docx")
